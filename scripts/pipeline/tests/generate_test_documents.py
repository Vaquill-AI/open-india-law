#!/usr/bin/env python3
"""
Generate test DOCX and Email files for parser testing.

Creates varied content to test:
- Legal document structures (pleadings, contracts, memos)
- Email formats (simple, with attachments, threads)
- Edge cases (Unicode, long text, special characters)
"""

import random
import string
from datetime import datetime, timedelta
from pathlib import Path
from email.message import EmailMessage
from email.utils import formatdate, make_msgid
import base64

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Will skip DOCX generation.")
    print("   Install with: pip install python-docx")

# Output directory
OUTPUT_DIR = Path(__file__).parent.parent.parent.parent / "data" / "document-tests"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Sample legal data
COURT_NAMES = [
    "Supreme Court of India", "Delhi High Court", "Bombay High Court",
    "Madras High Court", "Calcutta High Court"
]

CASE_TYPES = ["Civil Appeal", "Criminal Appeal", "Writ Petition", "Special Leave Petition"]

JUDGES = [
    "Hon'ble Justice D.Y. Chandrachud", "Hon'ble Justice B.R. Gavai",
    "Hon'ble Justice Surya Kant", "Hon'ble Justice Hima Kohli"
]

LEGAL_SECTIONS = [
    "Section 138 of the Negotiable Instruments Act, 1881",
    "Section 420 of the Indian Penal Code",
    "Article 226 of the Constitution of India",
    "Section 34 of the Arbitration and Conciliation Act, 1996"
]

LAWYER_NAMES = [
    "Mr. Harish Salve, Senior Advocate",
    "Mr. Mukul Rohatgi, Senior Advocate",
    "Ms. Indira Jaising, Senior Advocate",
    "Mr. Abhishek Manu Singhvi, Senior Advocate"
]


def random_date(start_year=2020, end_year=2024):
    """Generate random date."""
    start = datetime(start_year, 1, 1)
    end = datetime(end_year, 12, 31)
    delta = end - start
    random_days = random.randint(0, delta.days)
    return start + timedelta(days=random_days)


def random_legal_paragraph():
    """Generate a random legal-sounding paragraph."""
    templates = [
        "The petitioner submits that the impugned order dated {date} passed by the learned {court} is patently illegal and contrary to the settled principles of law.",
        "It is respectfully submitted that the respondent has failed to comply with the mandatory provisions of {section}, thereby vitiating the entire proceedings.",
        "The learned counsel for the appellant has relied upon the judgment of this Hon'ble Court in {case_name}, wherein it was held that procedural irregularities cannot be condoned.",
        "Upon consideration of the facts and circumstances of the case, and having heard the learned counsel for both parties, this Court is of the view that intervention is warranted.",
        "The respondent's contention that the petition is barred by limitation is devoid of merit, as the cause of action arose only on {date} when the impugned order was communicated.",
        "In light of the constitutional mandate under Article 21, the right to speedy trial is a fundamental right that cannot be denied to any citizen.",
        "The doctrine of legitimate expectation, as expounded by this Court in numerous judgments, mandates that the authorities act fairly and consistently.",
    ]

    return random.choice(templates).format(
        date=random_date().strftime("%d.%m.%Y"),
        court=random.choice(COURT_NAMES),
        section=random.choice(LEGAL_SECTIONS),
        case_name=f"XYZ v. State ({random.randint(2010, 2023)}) {random.randint(1, 12)} SCC {random.randint(1, 999)}"
    )


# ============================================================================
# DOCX GENERATORS
# ============================================================================

def generate_legal_pleading():
    """Generate a legal pleading document."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Title
    title = doc.add_heading("IN THE SUPREME COURT OF INDIA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case number
    doc.add_paragraph("CIVIL APPELLATE JURISDICTION", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"CIVIL APPEAL NO. {random.randint(1000, 9999)} OF {random.randint(2020, 2024)}", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Parties
    doc.add_paragraph("IN THE MATTER OF:")
    doc.add_paragraph(f"M/s ABC Industries Pvt. Ltd. ... APPELLANT")
    doc.add_paragraph("VERSUS")
    doc.add_paragraph(f"State of Maharashtra & Ors. ... RESPONDENTS")

    doc.add_paragraph()

    # Petition header
    doc.add_heading("SPECIAL LEAVE PETITION (CIVIL)", level=1)
    doc.add_paragraph("(Under Article 136 of the Constitution of India)")

    # Synopsis
    doc.add_heading("SYNOPSIS", level=2)
    for _ in range(3):
        doc.add_paragraph(random_legal_paragraph())

    # Grounds
    doc.add_heading("GROUNDS", level=2)
    for i in range(5):
        doc.add_paragraph(f"{i+1}. {random_legal_paragraph()}")

    # Prayer
    doc.add_heading("PRAYER", level=2)
    doc.add_paragraph("In view of the facts and circumstances stated above, it is most respectfully prayed that this Hon'ble Court may be pleased to:")
    doc.add_paragraph("(a) Grant leave to appeal against the impugned judgment;")
    doc.add_paragraph("(b) Set aside the impugned judgment and order;")
    doc.add_paragraph("(c) Pass such other order(s) as this Hon'ble Court may deem fit and proper.")

    # Signature
    doc.add_paragraph()
    doc.add_paragraph("Filed by:")
    doc.add_paragraph(random.choice(LAWYER_NAMES))
    doc.add_paragraph("Advocate for the Petitioner")
    doc.add_paragraph(f"Place: New Delhi")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y')}")

    filepath = OUTPUT_DIR / "legal_pleading.docx"
    doc.save(filepath)
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_legal_contract():
    """Generate a legal contract document."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Title
    title = doc.add_heading("SERVICE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Preamble
    doc.add_paragraph(
        f"This Service Agreement ('Agreement') is entered into as of {datetime.now().strftime('%B %d, %Y')} "
        f"('Effective Date') by and between:"
    )

    doc.add_paragraph("1. XYZ Technologies Private Limited, a company incorporated under the Companies Act, 2013, "
                     "having its registered office at 123 Tech Park, Bangalore - 560001 ('Service Provider'); AND")

    doc.add_paragraph("2. ABC Corporation Limited, a company incorporated under the Companies Act, 2013, "
                     "having its registered office at 456 Business Center, Mumbai - 400001 ('Client').")

    # Recitals
    doc.add_heading("RECITALS", level=1)
    doc.add_paragraph("WHEREAS, the Service Provider is engaged in the business of providing software development services;")
    doc.add_paragraph("WHEREAS, the Client desires to engage the Service Provider to provide certain services;")
    doc.add_paragraph("NOW, THEREFORE, in consideration of the mutual covenants contained herein, the parties agree as follows:")

    # Terms
    sections = [
        ("DEFINITIONS", [
            '"Confidential Information" means any information disclosed by either party...',
            '"Deliverables" means the work product to be delivered under this Agreement...',
            '"Services" means the services to be provided as described in Schedule A...'
        ]),
        ("SCOPE OF SERVICES", [
            "The Service Provider shall provide the services described in Schedule A.",
            "The Service Provider shall perform the services in a professional manner.",
            "Any changes to the scope shall be agreed in writing by both parties."
        ]),
        ("COMPENSATION", [
            "The Client shall pay the Service Provider as per the rates in Schedule B.",
            "Invoices shall be raised monthly and payable within 30 days.",
            "All amounts are exclusive of applicable taxes."
        ]),
        ("CONFIDENTIALITY", [
            "Each party shall maintain confidentiality of the other's Confidential Information.",
            "This obligation shall survive termination of this Agreement for 5 years.",
            "Disclosure may be made only with prior written consent or as required by law."
        ]),
        ("TERM AND TERMINATION", [
            "This Agreement shall commence on the Effective Date and continue for 12 months.",
            "Either party may terminate with 30 days written notice.",
            "Upon termination, all rights and licenses granted herein shall cease."
        ])
    ]

    for section_title, clauses in sections:
        doc.add_heading(section_title, level=1)
        for i, clause in enumerate(clauses):
            doc.add_paragraph(f"{i+1}. {clause}")

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")

    doc.add_paragraph()
    doc.add_paragraph("FOR XYZ TECHNOLOGIES PVT. LTD.")
    doc.add_paragraph("_____________________________")
    doc.add_paragraph("Authorized Signatory")

    doc.add_paragraph()
    doc.add_paragraph("FOR ABC CORPORATION LTD.")
    doc.add_paragraph("_____________________________")
    doc.add_paragraph("Authorized Signatory")

    filepath = OUTPUT_DIR / "legal_contract.docx"
    doc.save(filepath)
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_legal_memo():
    """Generate a legal memorandum."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # Header
    doc.add_heading("LEGAL MEMORANDUM", level=0)

    # Memo header
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "TO:"
    table.rows[0].cells[1].text = "Senior Partner"
    table.rows[1].cells[0].text = "FROM:"
    table.rows[1].cells[1].text = "Associate, Litigation Team"
    table.rows[2].cells[0].text = "DATE:"
    table.rows[2].cells[1].text = datetime.now().strftime("%B %d, %Y")
    table.rows[3].cells[0].text = "RE:"
    table.rows[3].cells[1].text = "Analysis of Limitation Period under Section 34 of Arbitration Act"

    doc.add_paragraph()

    # Question presented
    doc.add_heading("QUESTION PRESENTED", level=1)
    doc.add_paragraph(
        "Whether an application under Section 34 of the Arbitration and Conciliation Act, 1996 "
        "for setting aside an arbitral award is barred by limitation if filed beyond the prescribed "
        "period of three months from the date of receipt of the award?"
    )

    # Brief answer
    doc.add_heading("BRIEF ANSWER", level=1)
    doc.add_paragraph(
        "Yes, the application would be barred by limitation. However, the court has discretion "
        "to condone the delay up to 30 days beyond the initial three-month period under the "
        "proviso to Section 34(3), subject to sufficient cause being shown."
    )

    # Facts
    doc.add_heading("FACTS", level=1)
    for _ in range(3):
        doc.add_paragraph(random_legal_paragraph())

    # Discussion
    doc.add_heading("DISCUSSION", level=1)

    doc.add_heading("A. Statutory Framework", level=2)
    for _ in range(4):
        doc.add_paragraph(random_legal_paragraph())

    doc.add_heading("B. Judicial Interpretation", level=2)
    for _ in range(4):
        doc.add_paragraph(random_legal_paragraph())

    doc.add_heading("C. Application to Present Facts", level=2)
    for _ in range(3):
        doc.add_paragraph(random_legal_paragraph())

    # Conclusion
    doc.add_heading("CONCLUSION", level=1)
    doc.add_paragraph(
        "Based on the above analysis, it is concluded that the application under Section 34 "
        "must be filed within the prescribed limitation period. Any delay beyond 30 days from "
        "the expiry of the three-month period cannot be condoned under any circumstances."
    )

    filepath = OUTPUT_DIR / "legal_memo.docx"
    doc.save(filepath)
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_docx_with_tables():
    """Generate a DOCX with complex tables."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    doc.add_heading("CASE STATUS REPORT", level=0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_heading("Pending Cases Summary", level=1)

    # Create table
    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'

    # Header
    headers = ["Case No.", "Court", "Stage", "Next Date", "Status"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data
    stages = ["Arguments", "Evidence", "Final Hearing", "Judgment Reserved"]
    statuses = ["Active", "Adjourned", "Part-Heard", "Urgent"]

    for row_idx in range(1, 11):
        table.rows[row_idx].cells[0].text = f"CA/{random.randint(100, 999)}/2024"
        table.rows[row_idx].cells[1].text = random.choice(COURT_NAMES)
        table.rows[row_idx].cells[2].text = random.choice(stages)
        table.rows[row_idx].cells[3].text = random_date(2024, 2025).strftime("%d.%m.%Y")
        table.rows[row_idx].cells[4].text = random.choice(statuses)

    filepath = OUTPUT_DIR / "docx_with_tables.docx"
    doc.save(filepath)
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_docx_edge_cases():
    """Generate DOCX with edge cases."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    doc.add_heading("Edge Cases Test Document", level=0)

    # Unicode text
    doc.add_heading("Hindi Legal Text", level=1)
    doc.add_paragraph(
        "संविधान के अनुच्छेद 21 के तहत जीवन और व्यक्तिगत स्वतंत्रता का अधिकार एक मौलिक अधिकार है। "
        "इस अधिकार में त्वरित विचारण का अधिकार भी शामिल है।"
    )

    # Very long paragraph
    doc.add_heading("Long Paragraph", level=1)
    long_text = " ".join([random_legal_paragraph() for _ in range(10)])
    doc.add_paragraph(long_text)

    # Special characters
    doc.add_heading("Special Characters", level=1)
    doc.add_paragraph("Amounts: ₹1,00,00,000/- (Rupees One Crore Only)")
    doc.add_paragraph("Percentage: 18% GST @ ₹50,000/-")
    doc.add_paragraph('Quotes: "The judgment stated that..."')
    doc.add_paragraph("Section: § 34 of the Act")

    # Empty paragraphs
    doc.add_heading("Section with Gaps", level=1)
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("After empty paragraphs")

    filepath = OUTPUT_DIR / "docx_edge_cases.docx"
    doc.save(filepath)
    print(f"✓ Generated: {filepath}")
    return filepath


# ============================================================================
# EMAIL GENERATORS
# ============================================================================

def generate_simple_eml():
    """Generate a simple EML file."""
    msg = EmailMessage()
    msg['Subject'] = 'Re: Urgent - Status Update on Civil Appeal No. 1234/2024'
    msg['From'] = 'advocate@lawfirm.com'
    msg['To'] = 'client@company.com'
    msg['Cc'] = 'partner@lawfirm.com'
    msg['Date'] = formatdate(localtime=True)
    msg['Message-ID'] = make_msgid()

    body = """Dear Client,

I am writing to update you on the status of Civil Appeal No. 1234/2024 pending before the Supreme Court of India.

CURRENT STATUS:
The matter was listed for hearing today before the Hon'ble Justice D.Y. Chandrachud and Hon'ble Justice B.R. Gavai.

PROCEEDINGS:
1. The learned Senior Counsel for the Respondent sought time to file the reply.
2. The Court granted two weeks' time for filing the reply.
3. The matter has been adjourned to 15th February 2025.

NEXT STEPS:
1. We need to prepare rejoinder to the reply once filed.
2. I recommend we discuss the strategy before the next hearing.
3. Please confirm your availability for a meeting next week.

ACTION REQUIRED:
Please provide the following documents at the earliest:
- Audited financial statements for FY 2023-24
- Board resolution authorizing the appeal
- Any correspondence with the respondent

Please feel free to contact me if you have any queries.

Best regards,
Advocate Name
Partner, Litigation Department
Law Firm LLP
Phone: +91-11-12345678
Email: advocate@lawfirm.com
"""

    msg.set_content(body)

    filepath = OUTPUT_DIR / "simple_email.eml"
    with open(filepath, 'wb') as f:
        f.write(msg.as_bytes())
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_html_eml():
    """Generate an EML file with HTML content."""
    msg = EmailMessage()
    msg['Subject'] = 'Case Update: Important Legal Notice'
    msg['From'] = 'legal@corporation.com'
    msg['To'] = 'board@corporation.com'
    msg['Date'] = formatdate(localtime=True)
    msg['Message-ID'] = make_msgid()

    html_content = """
    <html>
    <head>
        <style>
            body { font-family: Arial, sans-serif; }
            .header { background-color: #003366; color: white; padding: 10px; }
            .content { padding: 20px; }
            table { border-collapse: collapse; width: 100%; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #003366; color: white; }
            .urgent { color: red; font-weight: bold; }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>Legal Department Update</h1>
        </div>
        <div class="content">
            <p>Dear Board Members,</p>

            <p class="urgent">URGENT: Action Required by Friday</p>

            <h2>Case Summary</h2>
            <table>
                <tr>
                    <th>Case Number</th>
                    <th>Court</th>
                    <th>Status</th>
                    <th>Next Date</th>
                </tr>
                <tr>
                    <td>WP 1234/2024</td>
                    <td>Delhi High Court</td>
                    <td>Arguments</td>
                    <td>20.01.2025</td>
                </tr>
                <tr>
                    <td>CA 5678/2024</td>
                    <td>Supreme Court</td>
                    <td>Listed</td>
                    <td>25.01.2025</td>
                </tr>
            </table>

            <h2>Key Points</h2>
            <ul>
                <li>Settlement discussions are ongoing</li>
                <li>Mediation scheduled for next month</li>
                <li>Board approval required for settlement authority</li>
            </ul>

            <p>Please review and confirm your availability.</p>

            <p>Best regards,<br>
            General Counsel</p>
        </div>
    </body>
    </html>
    """

    plain_content = """
    LEGAL DEPARTMENT UPDATE

    URGENT: Action Required by Friday

    Dear Board Members,

    Case Summary:
    - WP 1234/2024 | Delhi High Court | Arguments | 20.01.2025
    - CA 5678/2024 | Supreme Court | Listed | 25.01.2025

    Key Points:
    - Settlement discussions are ongoing
    - Mediation scheduled for next month
    - Board approval required for settlement authority

    Please review and confirm your availability.

    Best regards,
    General Counsel
    """

    msg.set_content(plain_content)
    msg.add_alternative(html_content, subtype='html')

    filepath = OUTPUT_DIR / "html_email.eml"
    with open(filepath, 'wb') as f:
        f.write(msg.as_bytes())
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_eml_with_attachment():
    """Generate an EML file with attachments."""
    msg = EmailMessage()
    msg['Subject'] = 'Documents for Review - Contract Agreement'
    msg['From'] = 'contracts@lawfirm.com'
    msg['To'] = 'client@business.com'
    msg['Date'] = formatdate(localtime=True)
    msg['Message-ID'] = make_msgid()

    body = """Dear Client,

Please find attached the following documents for your review:

1. Draft Service Agreement (contract_draft.txt)
2. Schedule A - Scope of Work (schedule_a.txt)

Please review and provide your comments by end of this week.

Best regards,
Contracts Team
"""

    msg.set_content(body)

    # Add text attachment 1
    attachment1_content = """
SERVICE AGREEMENT DRAFT
=======================

This is a draft service agreement between Party A and Party B.

TERMS:
1. Duration: 12 months
2. Value: INR 50,00,000
3. Payment: Monthly

[Further terms to be added]
"""
    msg.add_attachment(
        attachment1_content.encode('utf-8'),
        maintype='text',
        subtype='plain',
        filename='contract_draft.txt'
    )

    # Add text attachment 2
    attachment2_content = """
SCHEDULE A - SCOPE OF WORK
==========================

1. Phase 1: Requirements gathering
2. Phase 2: Development
3. Phase 3: Testing
4. Phase 4: Deployment

Timeline: 6 months from commencement
"""
    msg.add_attachment(
        attachment2_content.encode('utf-8'),
        maintype='text',
        subtype='plain',
        filename='schedule_a.txt'
    )

    filepath = OUTPUT_DIR / "email_with_attachments.eml"
    with open(filepath, 'wb') as f:
        f.write(msg.as_bytes())
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_thread_eml():
    """Generate an email thread (reply chain)."""
    # Create the threaded email with quoted content
    msg = EmailMessage()
    msg['Subject'] = 'Re: Re: Settlement Proposal - Case No. 789/2024'
    msg['From'] = 'lawyer@lawfirm.com'
    msg['To'] = 'opposing.counsel@otherfirm.com'
    msg['Date'] = formatdate(localtime=True)
    msg['Message-ID'] = make_msgid()
    msg['In-Reply-To'] = '<previous-message-id@otherfirm.com>'
    msg['References'] = '<original-message@lawfirm.com> <reply-1@otherfirm.com>'

    body = """Dear Counsel,

Thank you for your response. Our client is agreeable to the revised terms, subject to the following modifications:

1. Payment timeline extended to 90 days
2. Confidentiality clause to be mutual
3. Release to be limited to the current dispute

Please confirm if this is acceptable.

Best regards,
Advocate

---
On Mon, Jan 13, 2025 at 10:00 AM, opposing.counsel@otherfirm.com wrote:
>
> Dear Counsel,
>
> We have reviewed your settlement proposal. Our client is willing to consider
> the offer with the following conditions:
>
> 1. Full payment within 60 days
> 2. Confidentiality clause
> 3. Full and final release
>
> Please revert with your client's response.
>
> Best regards,
> Opposing Counsel
>
> ---
> On Fri, Jan 10, 2025 at 3:00 PM, lawyer@lawfirm.com wrote:
> >
> > Dear Counsel,
> >
> > Further to our discussion, we are proposing a settlement of INR 25,00,000
> > to resolve the matter amicably.
> >
> > Please take instructions from your client and revert.
> >
> > Best regards,
> > Advocate
"""

    msg.set_content(body)

    filepath = OUTPUT_DIR / "email_thread.eml"
    with open(filepath, 'wb') as f:
        f.write(msg.as_bytes())
    print(f"✓ Generated: {filepath}")
    return filepath


def generate_unicode_eml():
    """Generate an email with Unicode/Hindi content."""
    msg = EmailMessage()
    msg['Subject'] = 'कानूनी नोटिस - Legal Notice'
    msg['From'] = 'advocate@kanoon.in'
    msg['To'] = 'party@example.com'
    msg['Date'] = formatdate(localtime=True)
    msg['Message-ID'] = make_msgid()

    body = """सेवा में,
श्रीमान/श्रीमती,

विषय: कानूनी नोटिस

मेरे मुवक्किल की ओर से यह कानूनी नोटिस भेजा जा रहा है।

संविधान के अनुच्छेद 21 के तहत जीवन और व्यक्तिगत स्वतंत्रता का अधिकार एक मौलिक अधिकार है।

कृपया इस नोटिस की प्राप्ति के 15 दिनों के भीतर जवाब दें।

---

To,
Sir/Madam,

Subject: Legal Notice

This legal notice is being sent on behalf of my client.

The right to life and personal liberty under Article 21 of the Constitution is a fundamental right.

Please respond within 15 days of receipt of this notice.

धन्यवाद / Thank you,
अधिवक्ता / Advocate
"""

    msg.set_content(body)

    filepath = OUTPUT_DIR / "unicode_email.eml"
    with open(filepath, 'wb') as f:
        f.write(msg.as_bytes())
    print(f"✓ Generated: {filepath}")
    return filepath


def main():
    print("=" * 70)
    print("GENERATING TEST DOCUMENT FILES")
    print("=" * 70)
    print(f"Output directory: {OUTPUT_DIR}")
    print()

    files_generated = []

    # DOCX files
    print("\n--- DOCX Files ---")
    if DOCX_AVAILABLE:
        files_generated.append(generate_legal_pleading())
        files_generated.append(generate_legal_contract())
        files_generated.append(generate_legal_memo())
        files_generated.append(generate_docx_with_tables())
        files_generated.append(generate_docx_edge_cases())
    else:
        print("Skipping DOCX generation (python-docx not installed)")

    # Email files
    print("\n--- Email Files (EML) ---")
    files_generated.append(generate_simple_eml())
    files_generated.append(generate_html_eml())
    files_generated.append(generate_eml_with_attachment())
    files_generated.append(generate_thread_eml())
    files_generated.append(generate_unicode_eml())

    print("\n" + "=" * 70)
    print("GENERATION COMPLETE")
    print("=" * 70)
    valid_files = [f for f in files_generated if f]
    print(f"Total files generated: {len(valid_files)}")
    print(f"Location: {OUTPUT_DIR}")

    return valid_files


if __name__ == "__main__":
    main()
